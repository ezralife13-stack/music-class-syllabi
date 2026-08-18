from pathlib import Path
from math import cos, sin, pi
import textwrap

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "syllabus_assets"
OUT = ROOT / "Music_Curriculum_Master_Syllabus_PH.docx"
ASSET_DIR.mkdir(exist_ok=True)

# compact_reference_guide preset, resolved tokens
PAGE_W = 12240
PAGE_H = 15840
MARGIN = 1440
CONTENT_DXA = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FONT = "Calibri"
BODY_SIZE = 11
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2E7D7B"
GOLD = "7A5A00"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF7DB"
PALE_TEAL = "EAF5F4"
WHITE = "FFFFFF"
GRAY = "5B6573"
RED = "9B1C1C"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_font(run, name=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=TABLE_INDENT):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", total), ("w:tblInd", indent)):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa)-1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="B7C5D6", size="6", inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for edge in edges:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_cell_text(cell, text, bold=False, color=NAVY, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)


def add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_page_number(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    add_field(run, "PAGE")
    paragraph.add_run(" of ")
    run2 = paragraph.add_run()
    add_field(run2, "NUMPAGES")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = rgb(NAVY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Table Text" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(8.5)
        st.font.color.rgb = rgb(NAVY)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.08
    for name, fill, border in [("Callout", LIGHT_BLUE, BLUE), ("Teaching Note", PALE_GOLD, GOLD)]:
        if name not in [s.name for s in doc.styles]:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(9.5)
        st.font.color.rgb = rgb(NAVY)
        st.paragraph_format.left_indent = Inches(0.18)
        st.paragraph_format.right_indent = Inches(0.18)
        st.paragraph_format.space_before = Pt(6)
        st.paragraph_format.space_after = Pt(8)
        ppr = st.element.get_or_add_pPr()
        shd = ppr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            ppr.append(shd)
        shd.set(qn("w:fill"), fill)
        p_bdr = ppr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            ppr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:color"), border)
        left.set(qn("w:space"), "8")
        p_bdr.append(left)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs + [0]) + 1
    next_num = max(existing_num + [0]) + 1

    def make_abstract(abs_id, num_fmt, lvl_text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        txt = OxmlElement("w:lvlText")
        txt.set(qn("w:val"), lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, txt, jc, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rf = OxmlElement("w:rFonts")
            rf.set(qn("w:ascii"), font)
            rf.set(qn("w:hAnsi"), font)
            rpr.append(rf)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id, abs_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abs_id))
        num.append(aid)
        numbering.append(num)

    # OOXML requires all abstract numbering definitions before all concrete nums.
    make_abstract(next_abs, "bullet", "•", FONT)
    make_abstract(next_abs + 1, "decimal", "%1.")
    make_num(next_num, next_abs)
    bullet_id = next_num
    make_num(next_num + 1, next_abs + 1)
    return bullet_id, next_num + 1


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, nid])


doc = Document()
section = doc.sections[0]
doc.settings.odd_and_even_pages_header_footer = False
section.different_first_page_header_footer = False
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
configure_styles(doc)
BULLET_ID, NUMBER_ID = add_numbering(doc)


def add_bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, BULLET_ID)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, color=NAVY)
    else:
        r = p.add_run(text)
        set_font(r, color=NAVY)
    return p


def add_numbered(text):
    p = doc.add_paragraph()
    apply_num(p, NUMBER_ID)
    r = p.add_run(text)
    set_font(r, color=NAVY)
    return p


def add_labeled_paragraph(label, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(label + " ")
    set_font(r, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_font(r2, color=NAVY)
    return p


def add_callout(label, text, teaching=False):
    p = doc.add_paragraph(style="Teaching Note" if teaching else "Callout")
    r = p.add_run(label + " ")
    set_font(r, size=9.5, bold=True, color=GOLD if teaching else DARK_BLUE)
    r2 = p.add_run(text)
    set_font(r2, size=9.5, color=NAVY)
    return p


def add_header_footer(sec, left="MUSIC CURRICULUM", right="MASTER SYLLABUS"):
    hp = sec.header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    hr1 = hp.add_run(left)
    set_font(hr1, size=8, bold=True, color=GRAY)
    hr2 = hp.add_run("\t" + right)
    set_font(hr2, size=8, color=GRAY)
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "C7D3DF")
    pbdr.append(bottom)
    ppr.append(pbdr)
    fp = sec.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)
    for r in fp.runs:
        set_font(r, size=8, color=GRAY)


add_header_footer(section)


def page_break():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_title(title, subtitle=None, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(kicker.upper())
        set_font(r, size=9, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=24, bold=True, color=NAVY)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        sr = sp.add_run(subtitle)
        set_font(sr, size=11.5, italic=True, color=TEAL)


def add_heading(text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_table(headers, rows, widths, font_size=8.5, header_fill=LIGHT_BLUE, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=font_size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], BLUE)
    for ridx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            align = aligns[i] if aligns else (WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(cells[i], value, size=font_size, align=align)
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
    set_table_geometry(table, widths)
    set_table_borders(table)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return table


def add_metadata_table(items):
    rows = []
    for i in range(0, len(items), 2):
        left = items[i]
        right = items[i+1] if i+1 < len(items) else ("", "")
        rows.append((left[0], left[1], right[0], right[1]))
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, bold=idx in (0,2), color=DARK_BLUE if idx in (0,2) else NAVY, size=8.7)
            set_cell_shading(cells[idx], LIGHT_GRAY if idx in (0,2) else WHITE)
    set_table_geometry(table, [1350, 3330, 1350, 3330])
    set_table_borders(table, color="D1D9E2")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_course_outcomes(outcomes):
    rows = []
    for code, outcome, evidence in outcomes:
        rows.append((code, outcome, evidence))
    add_table(["Code", "By the end of the course, students can...", "Primary evidence"], rows,
              [900, 5820, 2640], font_size=8.6,
              aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])


def add_assessment_table(rows):
    add_table(["Assessment component", "Weight", "What it measures"], rows,
              [3600, 900, 4860], font_size=8.6,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])


def add_weekly_schedule(course_name, weeks):
    chunks = [weeks[0:6], weeks[6:12], weeks[12:18]]
    for ci, chunk in enumerate(chunks):
        page_break()
        add_section_title(f"{course_name}: Weekly Plan", f"Weeks {chunk[0][0]}-{chunk[-1][0]} | Assign exact textbook pages from the edition in use.", "Course Schedule")
        rows = []
        for w, focus, outcomes, activities, evidence in chunk:
            rows.append((w, focus, outcomes, activities, evidence))
        add_table(["Wk", "Focus / assigned reading", "Learning target", "Activity + visual aid", "Evidence"], rows,
                  [540, 2160, 2070, 2880, 1710], font_size=7.75,
                  aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
        add_callout("Instructor preparation:", "Select short examples that match student voices and instruments; use public-domain, licensed, teacher-composed, or institutionally cleared excerpts. Provide printed and projected versions.", teaching=True)


def add_signature_activities(course_name, activities):
    page_break()
    add_section_title(f"{course_name}: Activity Toolkit", "Ready-to-run protocols; scale the examples for mixed year levels.", "Teaching Toolkit")
    for title, time, setup, process, product in activities:
        add_heading(title, 2)
        add_labeled_paragraph("Time:", time, after=2)
        add_labeled_paragraph("Setup:", setup, after=2)
        add_labeled_paragraph("Process:", process, after=2)
        add_labeled_paragraph("Evidence:", product, after=5)


def add_rubric(title, criteria, note):
    add_heading(title, 2)
    rows = []
    for criterion, exemplary, proficient, developing, beginning in criteria:
        rows.append((criterion, exemplary, proficient, developing, beginning))
    add_table(["Criterion", "4 - Exemplary", "3 - Proficient", "2 - Developing", "1 - Beginning"], rows,
              [1400, 1990, 1990, 1990, 1990], font_size=7.35,
              aligns=[WD_ALIGN_PARAGRAPH.LEFT]*5)
    add_callout("Scoring note:", note)


# ---------- visual assets ----------
IMG_W = 1800
FONT_REG = Path("C:/Windows/Fonts/arial.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/arialbd.ttf")


def pil_font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REG
    return ImageFont.truetype(str(path), size=size)


def new_canvas(height):
    return Image.new("RGB", (IMG_W, height), "white"), ImageDraw.Draw(Image.new("RGB", (1, 1)))


def canvas(height):
    img = Image.new("RGB", (IMG_W, height), "white")
    return img, ImageDraw.Draw(img)


def centered(draw, xy, text, font, fill, spacing=4):
    draw.multiline_text(xy, text, font=font, fill="#"+fill, anchor="mm", align="center", spacing=spacing)


def arrow(draw, start, end, fill=DARK_BLUE, width=7):
    draw.line([start, end], fill="#"+fill, width=width)
    x1, y1 = start; x2, y2 = end
    ang = __import__("math").atan2(y2-y1, x2-x1)
    length = 24
    pts = [(x2, y2),
           (x2-length*cos(ang-pi/6), y2-length*sin(ang-pi/6)),
           (x2-length*cos(ang+pi/6), y2-length*sin(ang+pi/6))]
    draw.polygon(pts, fill="#"+fill)


def save_img(img, name):
    path = ASSET_DIR / name
    img.save(path, format="PNG", optimize=True)
    return path


def visual_learning_pathway():
    img, d = canvas(720)
    centered(d, (900, 75), "A Connected Musicianship Sequence", pil_font(50, True), NAVY)
    items = [
        (90, "MUSIC\nFUNDAMENTALS", "Read and write\nmusical language", LIGHT_BLUE),
        (520, "MUSIC THEORY I", "Build diatonic\nharmony", "DDEAF3"),
        (950, "MUSIC THEORY II", "Explain chromatic\nfunction", "D7E9E8"),
        (1380, "SIGHT SINGING", "Hear, prepare,\nand perform", "FFF1C2"),
    ]
    for i, (x, title, sub, fill) in enumerate(items):
        d.rounded_rectangle([x, 170, x+330, 510], radius=28, fill="#"+fill, outline="#"+BLUE, width=5)
        centered(d, (x+165, 280), title, pil_font(31, True), NAVY, spacing=8)
        centered(d, (x+165, 420), sub, pil_font(25), GRAY, spacing=6)
        if i < 3:
            arrow(d, (x+340, 340), (x+410, 340), TEAL, 6)
    centered(d, (900, 625), "Theory explains what the ear and voice experience; singing tests whether symbols have become sound.", pil_font(25), DARK_BLUE)
    return save_img(img, "V1_learning_pathway.png")


def visual_rhythm_tree():
    img, d = canvas(1050)
    centered(d, (900, 60), "Rhythmic Value Tree", pil_font(50, True), NAVY)
    levels = [
        ([(900, 175, "WHOLE\n4 quarter-note beats")], LIGHT_BLUE, NAVY, 280, 110, 27),
        ([(550, 410, "HALF\n2 beats"), (1250, 410, "HALF\n2 beats")], "DDEAF3", BLUE, 230, 100, 26),
        ([(300, 650, "QUARTER\n1 beat"),(700, 650, "QUARTER\n1 beat"),(1100, 650, "QUARTER\n1 beat"),(1500, 650, "QUARTER\n1 beat")], PALE_TEAL, TEAL, 210, 96, 23),
        ([(150+i*215, 895, "EIGHTH\n1/2 beat") for i in range(8)], PALE_GOLD, GOLD, 180, 88, 20),
    ]
    for nodes, fill, edge, bw, bh, fs in levels:
        for x,y,label in nodes:
            d.rounded_rectangle([x-bw//2,y-bh//2,x+bw//2,y+bh//2],radius=20,fill="#"+fill,outline="#"+edge,width=4)
            centered(d,(x,y),label,pil_font(fs,True),NAVY,spacing=3)
    for p, children in [((900,230),[(550,360),(1250,360)]),((550,460),[(300,600),(700,600)]),((1250,460),[(1100,600),(1500,600)])]:
        for c in children: d.line([p,c],fill="#"+GRAY,width=4)
    for x in [300,700,1100,1500]:
        for cx in [x-107,x+107]: d.line([(x,700),(cx,850)],fill="#"+GRAY,width=3)
    return save_img(img,"V2_rhythm_tree.png")


def visual_circle_fifths():
    img, d = canvas(1500)
    majors = ["C","G","D","A","E","B","F#/Gb","Db","Ab","Eb","Bb","F"]
    minors = ["Am","Em","Bm","F#m","C#m","G#m","D#m/Ebm","Bbm","Fm","Cm","Gm","Dm"]
    cx,cy=900,720
    for i,(maj,minr) in enumerate(zip(majors,minors)):
        angle=pi/2-i*2*pi/12
        x=int(cx+520*cos(angle)); y=int(cy-520*sin(angle))
        xi=int(cx+335*cos(angle)); yi=int(cy-335*sin(angle))
        d.ellipse([x-70,y-70,x+70,y+70],fill="#"+LIGHT_BLUE,outline="#"+BLUE,width=4)
        centered(d,(x,y),maj,pil_font(27,True),NAVY)
        d.ellipse([xi-58,yi-58,xi+58,yi+58],fill="#"+PALE_TEAL,outline="#"+TEAL,width=3)
        centered(d,(xi,yi),minr,pil_font(20),NAVY)
    d.ellipse([cx-190,cy-190,cx+190,cy+190],fill="#"+PALE_GOLD,outline="#"+GOLD,width=5)
    centered(d,(cx,cy-45),"CIRCLE OF",pil_font(32,True),NAVY)
    centered(d,(cx,cy+15),"FIFTHS",pil_font(47,True),GOLD)
    centered(d,(cx,cy+95),"outer: major\ninner: relative minor",pil_font(22),GRAY)
    centered(d,(480,1410),"Counterclockwise: add flats",pil_font(24),DARK_BLUE)
    centered(d,(1320,1410),"Clockwise: add sharps",pil_font(24),DARK_BLUE)
    return save_img(img,"V3_circle_of_fifths.png")


def visual_interval_flow():
    img,d=canvas(980)
    centered(d,(900,60),"Interval Identification Flow",pil_font(50,True),NAVY)
    nodes=[(900,185,580,120,"1. COUNT\nletter names / staff positions",LIGHT_BLUE,BLUE),
           (430,405,500,120,"2. FIND\nmajor or perfect reference",LIGHT_BLUE,BLUE),
           (1370,405,500,120,"3. COMPARE\nactual semitone distance",LIGHT_BLUE,BLUE),
           (430,650,460,118,"Perfect family\n1, 4, 5, 8",PALE_TEAL,TEAL),
           (1370,650,460,118,"Major/minor family\n2, 3, 6, 7",PALE_TEAL,TEAL),
           (900,875,600,120,"4. NAME THE QUALITY\nP / M / m / A / d",PALE_GOLD,GOLD)]
    for x,y,w,h,t,fill,edge in nodes:
        d.rounded_rectangle([x-w//2,y-h//2,x+w//2,y+h//2],radius=24,fill="#"+fill,outline="#"+edge,width=4)
        centered(d,(x,y),t,pil_font(27,True if t.startswith(("1","4")) else False),NAVY,spacing=5)
    for a,b in [((820,245),(570,340)),((980,245),(1230,340)),((430,465),(430,585)),((1370,465),(1370,585)),((560,710),(770,820)),((1240,710),(1030,820))]: arrow(d,a,b,GRAY,5)
    return save_img(img,"V4_interval_flow.png")


def visual_harmony_map():
    img,d=canvas(860)
    centered(d,(900,60),"Diatonic Harmonic Function Map",pil_font(50,True),NAVY)
    groups=[(275,"TONIC","I   •   vi   •   iii",LIGHT_BLUE,BLUE),(900,"PREDOMINANT","ii   •   IV",PALE_TEAL,TEAL),(1525,"DOMINANT","V   •   vii°",PALE_GOLD,GOLD)]
    for x,title,chords,fill,edge in groups:
        d.rounded_rectangle([x-220,220,x+220,580],radius=30,fill="#"+fill,outline="#"+edge,width=5)
        centered(d,(x,320),title,pil_font(32,True),NAVY)
        centered(d,(x,415),chords,pil_font(31),DARK_BLUE)
        centered(d,(x,510),"function, not a fixed order",pil_font(18),GRAY)
    arrow(d,(500,400),(675,400),GRAY,6); arrow(d,(1125,400),(1300,400),GRAY,6)
    d.arc([260,500,1540,810],0,180,fill="#"+RED,width=7); arrow(d,(285,655),(270,630),RED,7)
    centered(d,(900,790),"Most phrases move away from tonic, intensify toward dominant, and return to tonic.",pil_font(24),DARK_BLUE)
    return save_img(img,"V5_harmony_map.png")


def visual_voice_leading():
    img,d=canvas(980)
    centered(d,(900,60),"SATB Voice-Leading Checklist",pil_font(50,True),NAVY)
    items=[("1","SPELL","Correct chord members and inversion"),("2","RANGE","Keep each voice in a singable register"),("3","CONNECT","Retain common tones; move other voices smoothly"),("4","RESOLVE","Treat tendency tones and sevenths correctly"),("5","CHECK","No parallels, crossing, overlap, or hidden errors"),("6","SING","Hear every line, not only the vertical chord")]
    for i,(num,title,desc) in enumerate(items):
        y=185+i*120
        d.ellipse([105,y-35,175,y+35],fill="#"+(BLUE if i<3 else TEAL))
        centered(d,(140,y),num,pil_font(25,True),WHITE)
        d.text((215,y-22),title,font=pil_font(29,True),fill="#"+NAVY)
        d.text((520,y-20),desc,font=pil_font(27),fill="#"+GRAY)
        d.line([(215,y+46),(1650,y+46)],fill="#D9E1EA",width=3)
    centered(d,(900,925),"Write → play → sing → revise",pil_font(31,True),GOLD)
    return save_img(img,"V6_voice_leading.png")


def visual_sight_singing():
    img,d=canvas(840)
    centered(d,(900,60),"Sight-Singing Routine: Preview → Audiate → Perform",pil_font(46,True),NAVY)
    stages=[(330,"PREVIEW",["key / tonic","meter / pulse","range / leaps"],LIGHT_BLUE,BLUE),(900,"AUDIATE",["sing scale","tap rhythm","hear start + goal"],PALE_TEAL,TEAL),(1470,"PERFORM",["steady pulse","recover forward","reflect + retry"],PALE_GOLD,GOLD)]
    for i,(x,title,items,fill,edge) in enumerate(stages):
        d.ellipse([x-190,200,x+190,580],fill="#"+fill,outline="#"+edge,width=5)
        centered(d,(x,305),title,pil_font(32,True),NAVY)
        centered(d,(x,430),"\n".join(items),pil_font(25),GRAY,spacing=8)
        if i<2: arrow(d,(x+195,390),(stages[i+1][0]-195,390),DARK_BLUE,6)
    centered(d,(900,740),"Use movable-do solfège (default), scale-degree numbers, or a neutral syllable as assigned.",pil_font(24),DARK_BLUE)
    return save_img(img,"V7_sight_singing_routine.png")


VISUALS = [
    (visual_learning_pathway(), "V1. Connected Learning Pathway", "Use during orientation and advising to show how symbol, harmony, and audiation reinforce one another."),
    (visual_rhythm_tree(), "V2. Rhythmic Value Tree", "Use for duration equivalence, beat-unit changes, and subdivision drills."),
    (visual_circle_fifths(), "V3. Circle of Fifths", "Use for key signatures, relative keys, dominant relationships, and modulation planning."),
    (visual_interval_flow(), "V4. Interval Identification Flow", "Use whenever students confuse interval number and quality."),
    (visual_harmony_map(), "V5. Diatonic Harmonic Function Map", "Use for phrase analysis, progression building, and cadence diagnosis."),
    (visual_voice_leading(), "V6. SATB Voice-Leading Checklist", "Use as a desk reference during realization and correction labs."),
    (visual_sight_singing(), "V7. Sight-Singing Routine", "Project before every individual or ensemble reading."),
]


def add_picture_with_alt(path, width_inches, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    return p


# ---------- cover and front matter ----------
doc.add_paragraph().paragraph_format.space_after = Pt(70)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("COLLEGE MUSIC STUDIES")
set_font(r, size=10, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MASTER SYLLABUS PACKAGE")
set_font(r, size=29, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Music Fundamentals • Music Theory I • Music Theory II • Sight Singing")
set_font(r, size=14, bold=True, color=TEAL)
p.paragraph_format.space_after = Pt(20)
add_picture_with_alt(VISUALS[0][0], 6.15, "Four-course pathway from Music Fundamentals through Music Theory I and II to applied Sight Singing.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run("Designed for first- through fourth-year college students in the Philippines")
set_font(r, size=11, italic=True, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("18-week semester model | Outcomes-based | Mixed-experience friendly")
set_font(r, size=9.5, color=DARK_BLUE)

page_break()
add_section_title("Course Setup Sheet", "Complete once per course section before distributing to students.", "Editable Front Matter")
add_metadata_table([
    ("Institution", "[University / College]"), ("College / Department", "[Unit]"),
    ("Academic year", "[20__-20__]"), ("Semester", "[First / Second / Summer]"),
    ("Course code", "[Code]"), ("Section", "[Section]"),
    ("Instructor", "[Name and credentials]"), ("Email", "[Official email]"),
    ("Class schedule", "[Days / time]"), ("Room / platform", "[Room / LMS]"),
    ("Consultation", "[Schedule / booking process]"), ("Units / contact hours", "[Per institutional policy]"),
])
add_callout("Recommended scheduling assumption:", "Music Fundamentals, Theory I, and Theory II are planned for three 60-minute meetings or equivalent each week. Sight Singing is planned for two 60-minute laboratory meetings. Adjust units and hours to institutional rules.", teaching=True)
add_heading("Primary learning resources", 1)
for item in [
    "Surmani, Andrew; Surmani, Karen Farnum; and Manus, Morton. Alfred's Essentials of Music Theory. Use the edition currently held by the instructor for Music Fundamentals.",
    "Benner. Theory for Piano Students. G. Schirmer (publisher). Use the instructor's volume/edition for Music Theory I and II.",
    "Snyder, Audrey. The Sight Singer, Volume 1. Use the instructor's edition for graded reading practice.",
    "Instructor-prepared keyboard examples, staff paper, recordings, and short excerpts that are public-domain, licensed, teacher-composed, or institutionally cleared.",
]: add_bullet(item)
add_callout("Edition note:", "The weekly plans name topics instead of page numbers because editions and volumes differ. Enter the exact unit/page references before each term.")

page_break()
add_section_title("How to Use This Package", "A shared framework followed by four printable course sections and classroom appendices.", "Instructor Guide")
for n, text in enumerate([
    "Complete the Course Setup Sheet and copy the relevant details into each course section.",
    "Keep the sequence and outcomes; adjust repertoire, pacing, and textbook pages after the diagnostic.",
    "Use the activity toolkit for active practice. Each major concept should be heard, sung, played, written, and explained whenever possible.",
    "Project or print visual aids V1-V7. They are original summary graphics and do not reproduce textbook pages.",
    "Apply the included rubrics, then use the university's official grading conversion and attendance rules.",
]): add_numbered(text)
add_heading("Shared teaching principles", 1)
for item in [
    "Sound before symbol: begin with listening, tapping, singing, or keyboard exploration before formal labeling.",
    "Spiral review: retrieve earlier skills in every meeting instead of isolating them in one unit.",
    "Multiple representations: connect staff notation, keyboard geography, solfège/scale degree, Roman numeral, and sound.",
    "Mixed-level access: offer a core task plus support and extension so first- and fourth-year students can work meaningfully together.",
    "Philippine musical context: include responsibly attributed examples from Filipino art, folk, sacred, popular, and contemporary traditions when suitable; discuss context rather than treating repertoire as decoration.",
    "Feedback cycle: attempt, diagnose, coach, retry. Improvement after feedback is part of the learning evidence.",
]: add_bullet(item)
page_break()
add_heading("Minimum materials", 1)
add_metadata_table([
    ("Students", "Textbook access, pencil, eraser, staff paper, folder, headphones, phone/tablet when allowed"),
    ("Room", "Piano/keyboard, board, projector, speakers, metronome, moveable seating"),
    ("Instructor", "Prepared excerpts, answer keys, visual aids, LMS folders, accessible digital copies"),
    ("Optional", "MIDI keyboards, notation software, tuning app, document camera"),
])

add_section_title("Shared Policies and Learning Support", "Replace bracketed language with official institutional wording before release.", "Student-Facing Core")
policies = [
    ("Attendance and participation", "Regular attendance is essential because musicianship is cumulative and performance-based. Follow the institution's official absence and tardiness policy. Students should communicate documented or extended absences promptly and complete an agreed recovery plan."),
    ("Preparation", "Bring assigned materials and complete short pre-class practice. Written work must be legible; performed work should show evidence of a deliberate preview and practice process."),
    ("Late work and make-up assessment", "Use the institution's deadlines and make-up rules. Performance make-ups require comparable difficulty and may be scheduled individually or by recording when authorized."),
    ("Academic integrity", "Submit your own analysis, notation, singing, and composition. Collaboration is allowed only when stated. Cite borrowed musical examples and acknowledge assistance."),
    ("Generative AI and digital tools", "Notation, tuning, accompaniment, and AI tools may be used only when the task explicitly allows them. Students must disclose the tool and describe what it contributed. Tools may not replace a live skill check, closed assessment, or required personal analysis."),
    ("Accessibility and vocal care", "Students may request reasonable access arrangements through the appropriate university process. Sight-singing can be transposed to a healthy range without changing difficulty. No student is required to sing through pain or illness."),
    ("Respectful learning environment", "Use constructive language, honor names and identities, and discuss musical traditions with proper attribution and cultural care. Performance feedback addresses observable musical choices, not personal worth."),
    ("Privacy and recordings", "Record classmates only with explicit permission and follow institutional policy and applicable privacy requirements. Assessment recordings are for instructional use unless additional consent is obtained."),
]
for title, body in policies:
    if title == "Accessibility and vocal care":
        page_break()
        add_section_title("Shared Policies and Learning Support", "Accessibility, respectful participation, privacy, and institutional authority.", "Student-Facing Core continued")
    add_heading(title, 2)
    doc.add_paragraph(body)
add_callout("Institutional authority:", "Official university, college, department, student handbook, accessibility, data privacy, safety, grading, and academic integrity policies supersede this adaptable template.", teaching=True)

page_break()
add_section_title("Curriculum Alignment at a Glance", "The sequence develops literacy, analysis, creation, and independent musicianship.", "Program Map")
rows = [
    ("Music Fundamentals", "Notation, rhythm, scales, keys, intervals, triads, basic form", "Read • write • tap • play", "Diagnostic portfolio + musicianship project"),
    ("Music Theory I", "Diatonic harmony, Roman numerals, SATB voice leading, cadences, phrase structure", "Analyze • realize • compose • revise", "Harmonization portfolio + analysis"),
    ("Music Theory II", "Seventh chords, tonicization, modulation, chromatic predominant harmony, mixture, form", "Explain • transform • compose • defend", "Chromatic analysis-composition project"),
    ("Sight Singing", "Pulse, rhythm, solfège/scale degree, interval and tonal reading, independent recovery", "Preview • audiate • sing • reflect", "Progress portfolio + final prepared/unprepared readings"),
]
add_table(["Course", "Core knowledge", "Core practice", "Culminating evidence"], rows,
          [1500, 3150, 1830, 2880], font_size=8.2,
          aligns=[WD_ALIGN_PARAGRAPH.LEFT]*4)
add_heading("Shared program-level outcomes", 1)
for item in [
    "PLO 1 - Decode and produce standard music notation accurately.",
    "PLO 2 - Hear, perform, and explain rhythmic, melodic, and harmonic relationships.",
    "PLO 3 - Analyze tonal music using appropriate terminology and notation.",
    "PLO 4 - Create short, stylistically coherent musical work and revise it from feedback.",
    "PLO 5 - Practice independently, collaborate respectfully, and communicate musical decisions clearly.",
]: add_bullet(item)
add_heading("Recommended diagnostic in Week 1", 1)
add_metadata_table([
    ("Written (20 min)", "Staff reading, rhythm, key signatures, intervals, triad spelling"),
    ("Aural/performance (10 min)", "Pulse echo, pitch matching, short tonal/rhythmic reading"),
    ("Self-report (5 min)", "Training history, instrument/voice, notation confidence, access needs"),
    ("Use", "Place supports, form mixed-level pairs, set practice targets; do not use as a punitive grade"),
])


# ---------- course data ----------
courses = []

fund_outcomes = [
    ("MF1", "read treble and bass clefs and orient pitches on the grand staff and keyboard", "notation checks; keyboard lab"),
    ("MF2", "notate and perform common note/rest values in simple and compound meter", "rhythm labs; quizzes"),
    ("MF3", "construct major and minor scales and identify key signatures", "scale/key portfolio"),
    ("MF4", "identify and construct diatonic intervals and basic triads", "written and aural checks"),
    ("MF5", "recognize basic harmonic function, cadences, texture, phrase, and form", "listening map; analysis"),
    ("MF6", "apply vocabulary while listening, singing, playing, and discussing music", "oral explanation; activities"),
    ("MF7", "create and revise a short notated melody with coherent rhythm and phrase shape", "capstone composition"),
    ("MF8", "use a planned, reflective practice process", "practice log and reflection"),
]
fund_assess = [
    ("Short quizzes and notation checks", "20%", "Accurate retrieval of symbols, terms, and constructions"),
    ("Workbook / problem sets", "20%", "Guided application and correction habits"),
    ("Active musicianship labs", "15%", "Tapping, singing, keyboard, and collaborative work"),
    ("Midterm written-practical assessment", "20%", "Weeks 1-8 integration"),
    ("Final musicianship project + exam", "20%", "Creation, notation, performance, and explanation"),
    ("Professional learning habits", "5%", "Preparation, reflection, respectful participation"),
]
fund_weeks = [
    ("1", "Orientation; diagnostic; sound/symbol pathway", "Establish baseline and personal targets", "Diagnostic stations; V1 pathway; name-the-sound game", "Diagnostic profile"),
    ("2", "Staff, clefs, note names, ledger lines", "Read/write pitches in treble and bass", "Human staff; whiteboard relay; keyboard mapping", "Exit ticket"),
    ("3", "Grand staff, octave registers, accidentals", "Connect notation to pitch location", "Pitch scavenger hunt; error detective", "Notation check 1"),
    ("4", "Beat, tempo, note/rest values", "Maintain pulse and compare durations", "Body percussion; V2 rhythm tree; rhythm dominoes", "Rhythm lab 1"),
    ("5", "Simple meter; barlines; ties; dots", "Notate and perform simple-meter patterns", "Measure builder; conductor circle", "Quiz 1"),
    ("6", "Compound meter; subdivision", "Distinguish beat and subdivision", "Meter switch; rhythm translation", "Problem set"),
    ("7", "Major scale; whole/half-step pattern", "Construct and perform major scales", "Keyboard scale lab; floor-step pattern", "Scale check"),
    ("8", "Key signatures; circle of fifths", "Name and write major keys", "V3 circle; key-signature relay; flash retrieval", "Quiz 2"),
    ("9", "Midterm review and assessment", "Integrate notation, rhythm, scales, keys", "Team review circuit; individual written-practical", "Midterm"),
    ("10", "Minor scales and relative/parallel keys", "Construct natural, harmonic, melodic minor", "Minor makeover; compare sound and spelling", "Scale portfolio"),
    ("11", "Intervals: number, quality, inversion", "Identify and construct intervals", "V4 flow; interval ladder; partner checks", "Notation check 2"),
    ("12", "Triads: quality, root, inversion", "Spell and identify four triad qualities", "Triad factory; keyboard verification", "Quiz 3"),
    ("13", "Diatonic triads and Roman numerals", "Label triads within a key", "Chord card sorting; V5 function preview", "Problem set"),
    ("14", "Cadences and phrase endings", "Recognize basic cadence types", "Cadence corners; listening vote", "Listening map"),
    ("15", "Melody: motive, sequence, contour, phrase", "Describe and shape coherent melody", "Melody surgery; contour drawing", "Capstone draft"),
    ("16", "Texture, dynamics, articulation, basic form", "Use expressive and formal vocabulary", "Listening gallery; form cards", "Analysis paragraph"),
    ("17", "Capstone workshop and peer feedback", "Revise notation and musical logic", "Gallery walk; perform-and-revise conference", "Final project"),
    ("18", "Final integration; reflection", "Demonstrate cumulative literacy", "Written-practical exam; portfolio conference", "Final exam + reflection"),
]
fund_acts = [
    ("Human Grand Staff", "20-30 minutes", "Tape five-line staves on the floor; prepare pitch cards.", "Teams place themselves or cards on called pitches, then translate to the keyboard. Add accidentals and ledger lines for extension.", "Photo/diagram plus an individual five-item transfer check."),
    ("Rhythm Construction Lab", "35 minutes", "Prepare value cards, meter cards, and blank measures.", "Groups build measures, perform them, exchange with another group, and repair any meter violations. Advanced students add ties, dots, and syncopation.", "Correct measure set and performed reading."),
    ("Key-Signature Relay", "20 minutes", "Post a blank circle of fifths and prepare key cards.", "Students place keys, signatures, relative minors, and dominant relationships. Require each placement to be explained aloud.", "Timed accuracy plus one written explanation."),
    ("One-Phrase Composition Studio", "Two class meetings", "Provide an eight-measure frame and clear constraints.", "Students draft, sing/play, peer-review, revise, and submit a clean score with a short rationale.", "Notated melody, performance, and revision note."),
]

theory1_outcomes = [
    ("T1.1", "spell, invert, and classify intervals, triads, and seventh chords used in diatonic contexts", "construction quizzes"),
    ("T1.2", "analyze diatonic harmony with key, Roman numeral, inversion figure, and cadence", "analysis portfolio"),
    ("T1.3", "write stylistically controlled SATB progressions with correct spacing and tendency-tone treatment", "voice-leading labs"),
    ("T1.4", "realize figured-bass or Roman-numeral prompts at keyboard or in notation", "practical checks"),
    ("T1.5", "harmonize a melody or bass using functional progression and cadential closure", "harmonization project"),
    ("T1.6", "identify phrase, motive, sequence, cadence, and basic formal units", "annotated score"),
    ("T1.7", "diagnose and revise harmonic or voice-leading errors", "chorale repair tasks"),
    ("T1.8", "explain an analytical choice in clear musical language", "oral/written defense"),
]
theory1_assess = [
    ("Quizzes and timed constructions", "15%", "Fluent spelling, labeling, and recognition"),
    ("Problem sets / workbook", "20%", "Sustained practice and corrected reasoning"),
    ("Analysis and keyboard labs", "15%", "Transfer among score, sound, and instrument"),
    ("Midterm examination", "20%", "Diatonic analysis and foundational voice leading"),
    ("Harmonization portfolio + final exam", "25%", "Integrated analysis, writing, and revision"),
    ("Professional learning habits", "5%", "Preparation, collaboration, reflection"),
]
theory1_weeks = [
    ("1", "Diagnostic; tonal review; interval fluency", "Repair prerequisite gaps", "Retrieval grid; interval speed rounds; V4", "Diagnostic plan"),
    ("2", "Triads, inversions, open/close position", "Spell and arrange triads accurately", "Triad voicing cards; keyboard check", "Quiz 1"),
    ("3", "Key, scale degree, Roman numeral analysis", "Label diatonic triads in major/minor", "Roman-numeral stations; V3", "Problem set"),
    ("4", "SATB ranges, spacing, doubling", "Write singable four-part textures", "V6 checklist; voice-range audit", "Voice-leading lab 1"),
    ("5", "Melodic motion; parallels; tendency tones", "Connect chords without prohibited motion", "Chorale repair shop; sing each line", "Quiz 2"),
    ("6", "Root-position progressions; harmonic function", "Build functional phrase models", "V5 function map; progression tiles", "Mini-composition"),
    ("7", "Cadences and phrase structure", "Identify and write cadence types", "Cadence hunt; phrase brackets", "Analysis 1"),
    ("8", "First-inversion triads and bass-line shaping", "Use I6/IV6/ii6 appropriately", "Bass-line makeover; keyboard compare", "Voice-leading lab 2"),
    ("9", "Midterm review and examination", "Integrate analysis and voice leading", "Peer teach-back; individual exam", "Midterm"),
    ("10", "Second-inversion triads: passing, pedal, cadential", "Classify and control 6/4 functions", "6/4 sorting; error detective", "Quiz 3"),
    ("11", "Dominant seventh chord; preparation/resolution", "Resolve chordal seventh and leading tone", "Resolution arrows; sing inner voices", "Problem set"),
    ("12", "Diatonic seventh chords", "Analyze and resolve common seventh chords", "Chord-quality matrix; keyboard loop", "Analysis 2"),
    ("13", "Non-chord tones I: passing, neighbor, suspension", "Label and compose basic embellishments", "NCT color-coding; suspension chain", "Quiz 4"),
    ("14", "Non-chord tones II: appoggiatura, escape, anticipation, pedal", "Distinguish embellishments in context", "Listening/score stations", "Annotated score"),
    ("15", "Melody harmonization planning", "Choose bass, function, cadence, rhythm", "Harmonization decision tree", "Portfolio draft 1"),
    ("16", "Figured bass / Roman-numeral realization", "Realize a controlled progression", "Keyboard-to-paper relay", "Practical check"),
    ("17", "Form: period, sentence, binary/ternary overview", "Relate harmony to phrase design", "Phrase map gallery; peer conference", "Portfolio final"),
    ("18", "Final examination and reflection", "Demonstrate cumulative diatonic fluency", "Written analysis/realization; oral defense", "Final exam"),
]
theory1_acts = [
    ("Chorale Repair Shop", "35-45 minutes", "Prepare a short teacher-composed SATB passage with 8-10 deliberate errors.", "Teams diagnose by category, propose a minimal repair, sing/play each line, and defend one revision. Extensions require two valid solutions.", "Annotated before/after score and defense."),
    ("Progression Tile Lab", "30 minutes", "Create cards for Roman numerals, inversions, bass notes, and functions.", "Students build a phrase from tonic through predominant and dominant, test at keyboard, then voice-lead the best solution.", "Progression map plus four-part realization."),
    ("Cadence Court", "25 minutes", "Prepare short cadence examples.", "One group argues for a cadence label using soprano/bass scale degrees and harmony; another group challenges the evidence.", "Individual cadence justification."),
    ("Harmonization Conference", "Two class meetings", "Provide a melody with phrase marks and cadence goals.", "Students plan functions, draft bass, realize inner voices, peer-audit with V6, and revise after keyboard/singing feedback.", "Portfolio score, analysis, and revision memo."),
]

theory2_outcomes = [
    ("T2.1", "analyze and resolve diatonic and chromatic seventh chords", "analysis and writing"),
    ("T2.2", "identify and use secondary dominants and leading-tone chords", "tonicization studies"),
    ("T2.3", "distinguish tonicization from modulation and map pivot-chord processes", "modulation analysis"),
    ("T2.4", "analyze and write modal mixture, Neapolitan, and augmented-sixth sonorities", "chromatic harmony labs"),
    ("T2.5", "interpret altered chords through voice leading and harmonic function", "oral/written defense"),
    ("T2.6", "relate chromatic harmony to phrase, sequence, and larger form", "formal analysis"),
    ("T2.7", "compose and revise a short tonal passage using controlled chromaticism", "capstone composition"),
    ("T2.8", "compare multiple plausible analyses using evidence", "analysis colloquium"),
]
theory2_assess = [
    ("Quizzes and chord constructions", "15%", "Fast, accurate chromatic spelling and labeling"),
    ("Problem sets / analytical annotations", "20%", "Consistent application and evidence"),
    ("Keyboard, listening, and composition labs", "15%", "Sound-symbol transfer"),
    ("Midterm examination", "20%", "Seventh chords, tonicization, modulation"),
    ("Chromatic analysis-composition capstone + final", "25%", "Integrated explanation, creation, and revision"),
    ("Professional learning habits", "5%", "Preparation, discussion, feedback practice"),
]
theory2_weeks = [
    ("1", "Diagnostic; review of diatonic function", "Restore Theory I fluency", "Function-map recall; chorale repair; V5/V6", "Diagnostic task"),
    ("2", "Diatonic seventh chords in context", "Analyze inversion and resolution", "Resolution web; keyboard cycles", "Quiz 1"),
    ("3", "Sequences and expanded progression", "Track pattern and function", "Sequence strip puzzle; bass map", "Analysis 1"),
    ("4", "Secondary dominants", "Spell and resolve V/x", "Target-tonic game; arrow analysis", "Problem set"),
    ("5", "Secondary leading-tone chords", "Spell and resolve vii°/x and vii°7/x", "Chromatic spelling relay", "Quiz 2"),
    ("6", "Tonicization vs modulation", "Use duration, cadence, and hierarchy as evidence", "Evidence sort; listening timeline", "Short response"),
    ("7", "Pivot-chord modulation", "Identify common chord and new-key confirmation", "V3 route map; pivot passport", "Modulation lab"),
    ("8", "Direct/common-tone modulation; phrase plan", "Compare modulatory techniques", "Before/after recomposition", "Analysis 2"),
    ("9", "Midterm review and examination", "Integrate tonicization and modulation", "Colloquium review; exam", "Midterm"),
    ("10", "Mode mixture / borrowed chords", "Explain mixture by scale-degree alteration", "Parallel-key overlay; color coding", "Quiz 3"),
    ("11", "Neapolitan chord", "Spell, voice-lead, and contextualize N6", "Predominant comparison lab", "Writing study"),
    ("12", "Augmented-sixth chords", "Distinguish It+6, Fr+6, Ger+6", "Interval-expansion model; resolution singing", "Problem set"),
    ("13", "Altered dominants and enharmonic reinterpretation", "Trace voice-leading purpose", "Enharmonic pivot puzzle", "Quiz 4"),
    ("14", "Chromatic sequences and linear harmony", "Separate local voice-leading from global function", "Reduction workshop", "Analysis 3"),
    ("15", "Phrase/form with chromatic harmony", "Relate harmonic events to formal function", "Form-harmony timeline", "Capstone proposal"),
    ("16", "Capstone composition lab", "Use chromaticism with structural purpose", "Compose-play-sing-revise cycle", "Draft + conference"),
    ("17", "Analysis colloquium and portfolio revision", "Defend an interpretation with evidence", "Poster session; peer questions", "Capstone final"),
    ("18", "Final examination and reflection", "Demonstrate cumulative analytical control", "Analysis/writing exam; reflection", "Final exam"),
]
theory2_acts = [
    ("Tonicization Target Game", "30 minutes", "Prepare destination-chord cards and blank staff paper.", "Students draw a diatonic target, construct its applied dominant or leading-tone chord, resolve it, and explain altered tones. Advanced students create a two-link chain.", "Correct spelling, resolution, and functional explanation."),
    ("Modulation Route Map", "45 minutes", "Provide a short teacher-composed score that changes key.", "Students mark old key, pivot, reinterpretation, dominant preparation, cadence, and new key; then compare alternate readings.", "Annotated score and evidence statement."),
    ("Chromatic Predominant Laboratory", "40 minutes", "Prepare parallel phrase frames ending on V.", "Groups realize N6, It+6, Fr+6, and Ger+6 versions, sing critical voice-leading lines, and compare color and spacing.", "Four realizations plus comparison paragraph."),
    ("Analysis-Composition Colloquium", "Two class meetings", "Students bring capstone score and one-page harmonic map.", "Presenters perform/play the passage, explain chromatic choices, answer evidence-based questions, and revise after feedback.", "Final score, analysis, presentation, revision memo."),
]

sight_outcomes = [
    ("SS1", "establish tonic, pulse, meter, and starting pitch before reading", "preview checks"),
    ("SS2", "perform simple- and compound-meter rhythms with steady pulse", "rhythm readings"),
    ("SS3", "sing stepwise and triadic tonal patterns in major and minor", "melodic readings"),
    ("SS4", "read diatonic leaps through audiation and reference tones", "individual checks"),
    ("SS5", "maintain intonation, rhythmic continuity, and phrase direction", "performance rubric"),
    ("SS6", "recover musically after an error without stopping the pulse", "one-take readings"),
    ("SS7", "perform independently and in ensemble using solfège, scale degree, or assigned syllable", "weekly portfolio"),
    ("SS8", "diagnose a reading and plan the next practice step", "reflection log"),
]
sight_assess = [
    ("Weekly individual sight-singing checks", "30%", "First-read preparation, continuity, pitch, and rhythm"),
    ("Rhythm reading and ensemble checks", "15%", "Pulse, subdivision, meter, and coordination"),
    ("Practice / reflection portfolio", "10%", "Independent diagnosis and deliberate practice"),
    ("Midterm prepared + unprepared readings", "20%", "Cumulative Weeks 1-8 fluency"),
    ("Final prepared + unprepared readings", "20%", "Independent, musical reading at course level"),
    ("Professional learning habits", "5%", "Preparation, vocal care, supportive ensemble conduct"),
]
sight_weeks = [
    ("1", "Diagnostic; healthy posture/breath; matching pitch", "Establish starting skills and safe range", "Call-and-response; V7 routine; range check", "Diagnostic recording"),
    ("2", "Pulse, beat division, simple meter", "Read basic rhythms without pitch", "Walk-tap-speak; conductor circle", "Rhythm check 1"),
    ("3", "Do-re-mi patterns; stepwise major", "Read conjunct melodies", "Tonal ladder; hand-sign echo", "Melodic check 1"),
    ("4", "Skips within tonic triad", "Audiate do-mi-sol relationships", "Triad arpeggio relay; anchor-tone map", "Portfolio 1"),
    ("5", "Simple-meter rests, ties, dotted values", "Keep pulse through silence and sustain", "Silent-beat circle; rhythm canon", "Rhythm check 2"),
    ("6", "Leaps of 3rd and 4th", "Prepare and recover diatonic leaps", "Interval frame; partner drone", "Melodic check 2"),
    ("7", "Phrases, breath, dynamics, contour", "Read musically beyond accuracy", "Contour tracing; phrase relay", "Expressive reading"),
    ("8", "Compound meter and 6/8 patterns", "Feel two-level beat structure", "Macrobeat/microbeat movement", "Rhythm check 3"),
    ("9", "Midterm prepared and unprepared readings", "Integrate preview, rhythm, pitch, recovery", "Individual assessment and reflection", "Midterm"),
    ("10", "Minor mode and la-based minor default", "Establish and read minor tonal patterns", "Parallel-mode echo; minor anchor tones", "Melodic check 3"),
    ("11", "Leaps of 5th and 6th; inversion links", "Use reference tones for wider leaps", "Leap ladder; sing-fill-sing", "Portfolio 2"),
    ("12", "Anacrusis, syncopation, two-part rhythm", "Maintain meter through displacement", "Rhythm ensemble; role rotation", "Rhythm check 4"),
    ("13", "Sequences and changed starting degrees", "Track repeated patterns from new pitch levels", "Sequence staircase; memory echo", "Melodic check 4"),
    ("14", "Chromatic inflections as introduced in text", "Read altered tones from tonal context", "Approach-resolve drill; keyboard drone", "Individual check"),
    ("15", "Two-part / partner sight singing", "Maintain independent line and ensemble pulse", "Duet carousel; listening partner rubric", "Ensemble check"),
    ("16", "Longer mixed reading; efficient preview", "Prioritize hazards under time limit", "30-second scan challenge; coached retry", "Mock final"),
    ("17", "Final preparation and peer coaching", "Apply personal practice plan", "Sight-singing circuit; portfolio conference", "Portfolio final"),
    ("18", "Final prepared and unprepared readings", "Demonstrate independent course-level reading", "Individual assessment; reflection", "Final performance"),
]
sight_acts = [
    ("Thirty-Second Scan", "10-15 minutes daily", "Project one short teacher-selected melody.", "Students silently identify key/tonic, meter, starting degree, rhythmic hazards, leaps, and phrase goal; then sing once without stopping and record one next step.", "Preview checklist and one-sentence reflection."),
    ("Recovery Drill", "15 minutes", "Use a melody with clear tonic/dominant anchors.", "The instructor deliberately signals an error point. Students keep pulse, aim for the next anchor, and rejoin. Debrief forward recovery rather than restarting.", "Continuity score and identified anchor."),
    ("Partner Drone and Leap Lab", "20-25 minutes", "Pairs, keyboard or tuning reference, interval cards.", "One student sustains tonic/dominant while the other sings target patterns; switch roles. Advanced students begin on non-tonic degrees.", "Accurate pattern set and peer feedback."),
    ("One-Take / Coached Retry Portfolio", "Weekly", "Recording device when authorized; otherwise live conference.", "Perform a first take, annotate errors, choose one strategy, practice briefly, then record a second take without editing.", "Both attempts plus concise diagnosis."),
]

courses = [
    ("Music Fundamentals", "A foundation course in music literacy linking notation, rhythm, scales, keys, intervals, triads, listening, keyboard geography, and basic form. Designed for students with varied prior training.", "None; placement or diagnostic information is used for support, not exclusion.", "Recommended: 3 units / 54 contact hours", fund_outcomes, fund_assess, fund_weeks, fund_acts),
    ("Music Theory I", "A first course in tonal analysis and writing emphasizing diatonic harmony, Roman numeral analysis, SATB voice leading, cadences, non-chord tones, phrase structure, and basic form.", "Music Fundamentals or equivalent placement; keyboard familiarity is helpful.", "Recommended: 3 units / 54 contact hours", theory1_outcomes, theory1_assess, theory1_weeks, theory1_acts),
    ("Music Theory II", "A continuation of tonal theory emphasizing seventh chords, tonicization, modulation, chromatic predominant harmony, mixture, altered harmony, formal function, analysis, and composition.", "Music Theory I or equivalent.", "Recommended: 3 units / 54 contact hours", theory2_outcomes, theory2_assess, theory2_weeks, theory2_acts),
    ("Sight Singing", "A progressive laboratory in rhythmic and melodic reading that develops audiation, tonal orientation, intonation, continuity, ensemble awareness, and reflective practice.", "Concurrent or prior Music Fundamentals recommended. Default system: movable-do with la-based minor; the instructor may substitute another consistent system.", "Recommended: 1 unit / 36 contact hours", sight_outcomes, sight_assess, sight_weeks, sight_acts),
]


for ci, (name, desc, prereq, units, outcomes, assessments, weeks, activities) in enumerate(courses):
    page_break()
    add_section_title(name, desc, f"Course {ci+1} of 4")
    add_metadata_table([
        ("Course code", "[Insert code]"), ("Units / hours", units),
        ("Prerequisite", prereq), ("Primary text", ["Alfred's Essentials of Music Theory", "Benner, Theory for Piano Students", "Benner, Theory for Piano Students", "Snyder, The Sight Singer, Vol. 1"][ci]),
        ("Delivery", "In-person / blended / online as approved"), ("Class section", "[Insert section]"),
    ])
    add_heading("Course learning outcomes", 1)
    add_course_outcomes(outcomes)
    page_break()
    add_section_title(f"{name}: Assessment Plan", "Weights, evidence, and the shared improvement cycle.", "Assessment and Grading")
    add_heading("Assessment and grading", 1)
    add_assessment_table(assessments)
    add_callout("Grade calculation:", "Calculate the weighted raw score shown above, then apply the official university grading/transmutation system. Publish the approved equivalent table in the LMS or course handout.")
    add_heading("Learning evidence cycle", 1)
    for step in ["Prepare: short reading/practice before class.", "Attempt: low-stakes first try.", "Diagnose: label the specific musical issue.", "Coach and revise: apply one targeted strategy.", "Demonstrate: submit or perform independently."]:
        add_bullet(step)
    add_weekly_schedule(name, weeks)
    add_signature_activities(name, activities)


# ---------- common rubrics ----------
page_break()
add_section_title("Assessment Rubrics", "Use the relevant rubric, add task-specific criteria if needed, and publish before assessment.", "Appendix A")
analysis_criteria = [
    ("Accuracy", "Labels, spelling, and notation are consistently correct.", "Minor errors do not change the main interpretation.", "Several errors weaken the interpretation.", "Frequent errors prevent a valid reading."),
    ("Reasoning", "Claims use precise, sufficient musical evidence.", "Most claims have relevant evidence.", "Evidence is partial or inconsistently linked.", "Claims are unsupported or unclear."),
    ("Completeness", "All required layers are addressed and connected.", "Required layers are substantially complete.", "Important layers are missing.", "Response is fragmentary."),
    ("Communication", "Notation and prose are clear, concise, and professional.", "Meaning is clear with small presentation issues.", "Presentation sometimes obscures meaning.", "Presentation is difficult to follow."),
]
add_rubric("Written Analysis / Theory Work", analysis_criteria, "Multiply criteria by the task weighting. For timed work, communicate any adjusted emphasis in advance.")

page_break()
add_section_title("Assessment Rubrics", "Performance and composition criteria.", "Appendix A continued")
composition_criteria = [
    ("Technical control", "Harmony, notation, and voice leading are controlled throughout.", "Mostly controlled with minor lapses.", "Inconsistent control; revision is needed.", "Technical issues obscure the musical plan."),
    ("Musical coherence", "Phrase, function, and form create a convincing whole.", "Plan is clear and mostly coherent.", "Some sections lack connection or direction.", "No sustained structural plan is evident."),
    ("Craft and purpose", "Choices are intentional and audibly effective.", "Most choices support the stated purpose.", "Purpose is unevenly realized.", "Choices seem arbitrary or copied without understanding."),
    ("Revision", "Feedback is evaluated and revisions materially improve the work.", "Relevant revisions are completed.", "Revision is superficial or incomplete.", "No meaningful revision is shown."),
]
add_rubric("Composition / Harmonization", composition_criteria, "Require a draft, feedback record, and final version so revision is visible.")
sight_criteria = [
    ("Preparation", "Efficiently identifies tonic, meter, hazards, and phrase plan.", "Identifies most essential features.", "Preview is incomplete or unfocused.", "Begins without an effective preview."),
    ("Rhythm", "Pulse, subdivision, and durations remain secure.", "Minor disruptions; pulse is recovered quickly.", "Several disruptions affect continuity.", "Pulse and meter are not sustained."),
    ("Pitch / intonation", "Tonal center and intervals remain accurate.", "Mostly accurate with quick recovery.", "Frequent errors weaken tonal orientation.", "Pitch relationships are largely inaccurate."),
    ("Continuity / expression", "Performs forward with phrasing, breath, and musical shape.", "Maintains continuity with basic expression.", "Stops/restarts or gives limited phrase direction.", "Cannot sustain the reading."),
]
add_rubric("Sight-Singing Performance", sight_criteria, "Suggested weighting: Preparation 15%, Rhythm 30%, Pitch/intonation 35%, Continuity/expression 20%. Transpose to a healthy range without penalty.")


# ---------- visual aids appendix ----------
for idx, (path, title, use) in enumerate(VISUALS[1:], start=2):
    page_break()
    add_section_title(title, use, "Appendix B • Printable Visual Aid")
    width = 5.9 if idx == 3 else 6.2
    add_picture_with_alt(path, width, title + ". " + use)
    add_heading("Three ways to use it", 2)
    if idx == 2:
        tips = ["Cover the lower levels and ask students to predict subdivisions.", "Change the beat unit and relabel the equivalences.", "Have groups build one measure using one path through the tree."]
    elif idx == 3:
        tips = ["Trace clockwise to add sharps and counterclockwise to add flats.", "Pair each major key with its relative minor.", "Draw a route from an original key to a possible destination key."]
    elif idx == 4:
        tips = ["Require the number before the quality.", "Use a keyboard only after students predict the reference interval.", "Invert the interval and explain how number and quality change."]
    elif idx == 5:
        tips = ["Sort real chords by function before naming exact numerals.", "Build two different predominant paths to the same cadence.", "Mark where a phrase departs from and returns to tonic."]
    elif idx == 6:
        tips = ["Audit one category at a time instead of scanning randomly.", "Sing each line after vertical errors are corrected.", "Ask students to justify the smallest effective revision."]
    else:
        tips = ["Use the same routine before every graded reading.", "Limit preview time gradually as fluency improves.", "After singing, choose only one priority for the coached retry."]
    for tip in tips: add_bullet(tip)
    add_callout("Accessibility:", "Provide the image digitally with the embedded description, read key relationships aloud, and allow students to recreate it in a tactile or high-contrast format as needed.")


# ---------- activity bank and planning ----------
page_break()
add_section_title("Reusable Activity Bank", "Short routines that can be inserted into any of the four courses.", "Appendix C")
bank = [
    ("Five-minute retrieval grid", "Individual recall, then pair correction; use yesterday/last week/last month items."),
    ("Error detective", "Students explain why a prepared notation, analysis, or performance is wrong and repair it minimally."),
    ("Hear-sing-play-write", "Present one concept in four representations; rotate the starting mode each week."),
    ("Whiteboard relay", "Teams solve one step each; a final verifier checks the entire chain, not just the answer."),
    ("Gallery walk", "Display analyses or scores; peers leave one evidence-based affirmation and one question."),
    ("Think-pair-perform", "Silent preparation, partner rehearsal, then public performance or explanation."),
    ("Confidence traffic light", "Students mark green/yellow/red privately; use results to form temporary support groups."),
    ("Exit ticket", "One construction, one explanation, and one next practice step."),
]
for title, desc in bank:
    add_heading(title, 2)
    doc.add_paragraph(desc)
page_break()
add_section_title("Mixed-Experience Differentiation", "Core, support, and extension options for the same learning target.", "Appendix C continued")
add_table(["Core task", "Support option", "Extension option"], [
    ("Read or analyze the assigned example", "Color-coded guide tones, reduced length, paired preview, keyboard reference", "Transpose, generate an alternate solution, or defend ambiguity"),
    ("Write a progression or melody", "Provide phrase frame, rhythm bank, or bass skeleton", "Add controlled chromaticism, secondary line, or formal expansion"),
    ("Sight-sing the assigned line", "Slower tempo, tonic drone, chunked rehearsal, comfortable transposition", "Reduced prep time, non-tonic start, partner line, or neutral syllable"),
], [2700, 3330, 3330], font_size=8.3, aligns=[WD_ALIGN_PARAGRAPH.LEFT]*3)

page_break()
add_section_title("Lesson Planning Template", "A repeatable 60-minute structure; extend proportionally for longer meetings.", "Appendix D")
lesson_rows = [
    ("0-5 min", "Retrieval / pulse warm-up", "One familiar skill, brisk feedback"),
    ("5-12 min", "Sound first", "Listen, sing, tap, or play the new idea"),
    ("12-25 min", "Model and label", "Connect sound to notation and vocabulary"),
    ("25-42 min", "Guided active task", "Pairs/stations; instructor checks misconceptions"),
    ("42-52 min", "Independent transfer", "Short written, keyboard, or performance problem"),
    ("52-58 min", "Feedback and retry", "One targeted correction; immediate second attempt"),
    ("58-60 min", "Exit ticket", "Evidence plus next practice step"),
]
add_table(["Time", "Learning phase", "Instructor / student action"], lesson_rows,
          [1200, 2700, 5460], font_size=8.7,
          aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
add_heading("Before class checklist", 1)
for item in ["Outcome and success criterion are visible.", "Examples fit the target skill and student range.", "Visual aid / board sequence is prepared.", "Support and extension versions are ready.", "Assessment evidence can be completed within the available time.", "Audio, keyboard, projector, and file access are tested."]:
    add_bullet(item)
add_heading("After class reflection", 1)
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(0)
spacer.paragraph_format.space_after = Pt(2)
for idx, q in enumerate(["What could most students do independently?", "Which misconception needs retrieval next meeting?", "Who needs support or extension?", "Did students connect sound, symbol, and explanation?", "What should be shortened, repeated, or replaced?"]):
    reflection_item = add_bullet(q)
    if idx == 0:
        reflection_item.paragraph_format.space_before = Pt(8)

page_break()
add_section_title("Instructor Readiness Checklist", "Finalize these items before releasing the syllabus.", "Appendix E")
for item in [
    "Insert official institution, department, course code, units, schedule, and faculty information.",
    "Confirm the academic calendar, holidays, examination week, and number of meetings.",
    "Insert exact textbook volume, edition, ISBN if required, and weekly page assignments.",
    "Replace template policy language with official institutional text and links.",
    "Confirm the official grading scale, transmutation, attendance, and make-up procedures.",
    "Choose the sight-singing tonal system and communicate it consistently.",
    "Choose assessment excerpts and verify copyright/licensing or institutional permission.",
    "Prepare accessible digital files and arrange accommodations through the proper process.",
    "Publish rubrics, due dates, submission formats, and recording/privacy expectations.",
    "Run the Week 1 diagnostic and revise pacing without lowering the stated outcomes.",
]: add_bullet("☐ " + item)
add_callout("Recommended first customization:", "Tell students exactly which textbook volume/pages correspond to each week. This package intentionally avoids edition-specific page references so it remains accurate across instructor copies.", teaching=True)
add_heading("Closing statement for students", 1)
doc.add_paragraph("Musicianship grows through frequent, focused attempts. Accuracy matters, but so do listening, recovery, explanation, revision, and respect for the people and traditions represented in the music. Come prepared to make sound, make decisions, and improve them.")


# document core properties and final options
props = doc.core_properties
props.title = "Master Syllabus Package: Music Fundamentals, Music Theory I-II, and Sight Singing"
props.subject = "College-level music curriculum for Philippine higher education"
props.author = "Course Instructor"
props.keywords = "music fundamentals, music theory, sight singing, syllabus, Philippines"

# Keep tables/figures readable and set update-fields-on-open.
settings = doc.settings.element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.save(OUT)
print(OUT)
