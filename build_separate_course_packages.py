from __future__ import annotations

import ast
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "build_syllabus.py"
ASSET_DIR = ROOT / "syllabus_assets"


def load_master_components():
    """Load only reusable helpers and course-data assignments from the master builder."""
    tree = ast.parse(SOURCE.read_text(encoding="utf-8"), filename=str(SOURCE))
    helper_names = {
        "rgb", "set_font", "set_cell_shading", "set_cell_margins",
        "set_repeat_table_header", "set_table_geometry", "set_table_borders",
        "set_keep_with_next", "set_cell_text", "add_field", "add_page_number",
        "configure_styles", "add_numbering", "apply_num", "add_bullet",
        "add_numbered", "add_labeled_paragraph", "add_callout",
        "add_header_footer", "page_break", "add_section_title", "add_heading",
        "add_table", "add_metadata_table", "add_course_outcomes",
        "add_assessment_table", "add_weekly_schedule", "add_signature_activities",
        "add_rubric",
    }
    data_names = {
        f"{prefix}_{suffix}"
        for prefix in ("fund", "theory1", "theory2", "sight")
        for suffix in ("outcomes", "assess", "weeks", "acts")
    } | {"courses"}

    function_nodes = [
        node for node in tree.body
        if isinstance(node, ast.FunctionDef) and node.name in helper_names
    ]
    data_nodes = []
    for node in tree.body:
        if not isinstance(node, ast.Assign):
            continue
        target_names = {t.id for t in node.targets if isinstance(t, ast.Name)}
        if target_names & data_names:
            data_nodes.append(node)

    exec(compile(ast.fix_missing_locations(ast.Module(function_nodes, type_ignores=[])), str(SOURCE), "exec"), globals())
    exec(compile(ast.fix_missing_locations(ast.Module(data_nodes, type_ignores=[])), str(SOURCE), "exec"), globals())


# Resolved compact_reference_guide design tokens used by the source helpers.
PAGE_W = 12240
PAGE_H = 15840
MARGIN = 1440
CONTENT_DXA = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FONT = "Calibri"
BODY_SIZE = 11
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2E7D7B"
GOLD = "7A5A00"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF7DB"
PALE_TEAL = "EAF5F4"
WHITE = "FFFFFF"
GRAY = "5B6573"
RED = "9B1C1C"

load_master_components()


COURSE_CONFIG = {
    "Music Fundamentals": {
        "filename": "Music_Fundamentals_Syllabus_and_Lesson_Plans.docx",
        "code": "[Insert course code]",
        "text": "Alfred's Essentials of Music Theory, Andrew Surmani, Karen Farnum Surmani, and Morton Manus",
        "meeting_model": "Suggested model: three 60-minute meetings weekly; adjust to the approved institutional schedule.",
        "meeting_names": ["Hear and discover", "Practice and apply", "Demonstrate and reflect"],
        "visuals": [
            ("V2_rhythm_tree.png", "Rhythm Value Tree", "A hierarchy showing relationships among common note values and subdivisions."),
            ("V3_circle_of_fifths.png", "Circle of Fifths", "A circle diagram linking key signatures, relative keys, and fifth relationships."),
            ("V4_interval_flow.png", "Interval Identification Flow", "A decision flow for counting interval number and determining quality."),
            ("V5_harmony_map.png", "Tonal Function Map", "A functional map connecting tonic, predominant, and dominant harmony."),
        ],
        "support": "Offer a labeled staff or keyboard map, color-coded scale degrees, a reduced item set, and a worked first example.",
        "extension": "Require transposition, a second valid solution, an original two-measure example, or a concise explanation of the governing rule.",
        "practice": "Complete the corresponding Alfred's Essentials of Music Theory unit(s), correct every missed item, and bring one question or self-composed example.",
        "rubric_title": "Music Literacy Task Rubric",
        "rubric_note": "Use for notation, rhythm, scales, intervals, and chord-construction tasks. Convert the four criteria to the task's announced point value.",
        "rubric": [
            ("Notation accuracy", "Symbols, placement, and spelling are consistently correct.", "Minor errors do not obscure meaning.", "Several errors show partial control.", "Frequent errors prevent a readable answer."),
            ("Concept application", "Chooses and applies the correct process independently.", "Process is sound with small slips.", "Needs prompts to select or finish the process.", "Uses an unsuitable or incomplete process."),
            ("Aural-symbol link", "Connects sound, notation, and terminology precisely.", "Connection is generally accurate.", "Connection is inconsistent.", "Connection is not yet demonstrated."),
            ("Revision", "Diagnoses and corrects errors with a clear reason.", "Corrects most errors after feedback.", "Corrects some errors with support.", "Does not yet use feedback effectively."),
        ],
    },
    "Music Theory I": {
        "filename": "Music_Theory_I_Syllabus_and_Lesson_Plans.docx",
        "code": "[Insert course code]",
        "text": "Theory for Piano Students, G. Schirmer (Benner)",
        "meeting_model": "Suggested model: three 60-minute meetings weekly; adjust to the approved institutional schedule.",
        "meeting_names": ["Hear and analyze", "Write and test", "Demonstrate and revise"],
        "visuals": [
            ("V4_interval_flow.png", "Interval Identification Flow", "A decision flow supporting interval spelling and harmonic construction."),
            ("V5_harmony_map.png", "Tonal Function Map", "A functional map connecting tonic, predominant, and dominant harmony."),
            ("V6_voice_leading.png", "SATB Voice-Leading Checklist", "A checklist for ranges, spacing, doubling, tendency tones, and parallels."),
        ],
        "support": "Provide a Roman-numeral/function cue, bass skeleton, limited chord vocabulary, keyboard model, and a singable inner-line check.",
        "extension": "Ask for an alternate voicing or harmonization, keyboard realization in a new key, or an evidence-based defense of two plausible analyses.",
        "practice": "Complete the corresponding Benner section, label function and inversion, play or sing critical lines, then submit corrected work with one-sentence error diagnoses.",
        "rubric_title": "Diatonic Analysis and Voice-Leading Rubric",
        "rubric_note": "Apply to written analysis, SATB work, figured-bass or Roman-numeral realization, and harmonization projects.",
        "rubric": [
            ("Harmonic accuracy", "Spelling, inversion, function, and cadence labels are consistently correct.", "Minor errors do not change the analysis.", "Several errors weaken the analysis.", "Frequent errors prevent a valid reading."),
            ("Voice leading", "Ranges, spacing, doubling, motion, and resolutions are controlled.", "Mostly controlled with isolated faults.", "Recurring faults require revision.", "Writing is not yet singable or functional."),
            ("Musical evidence", "Every claim is supported by precise score or aural evidence.", "Most claims use relevant evidence.", "Evidence is partial or loosely connected.", "Claims are unsupported."),
            ("Revision and clarity", "Notation is clear and revisions solve the diagnosed issue.", "Work is readable and revisions mostly succeed.", "Presentation or revisions remain inconsistent.", "Work is incomplete or difficult to follow."),
        ],
    },
    "Music Theory II": {
        "filename": "Music_Theory_II_Syllabus_and_Lesson_Plans.docx",
        "code": "[Insert course code]",
        "text": "Theory for Piano Students, G. Schirmer (Benner)",
        "meeting_model": "Suggested model: three 60-minute meetings weekly; adjust to the approved institutional schedule.",
        "meeting_names": ["Hear and map", "Analyze and realize", "Defend and revise"],
        "visuals": [
            ("V3_circle_of_fifths.png", "Circle of Fifths", "A key-relationship map for tonicization and modulation routes."),
            ("V5_harmony_map.png", "Tonal Function Map", "A function map extended through applied and chromatic harmony."),
            ("V6_voice_leading.png", "Chromatic Voice-Leading Checklist", "A checklist for altered tones, tendency tones, spacing, and resolution."),
        ],
        "support": "Use a chromatic-spelling scaffold, circle-of-fifths route map, reduction template, keyboard drone, and a limited choice of destination keys.",
        "extension": "Require enharmonic reinterpretation, an alternate modulation or analysis, recomposition with a different chromatic predominant, or a formal-function defense.",
        "practice": "Complete the corresponding Benner section, annotate altered tones and resolutions, test the progression at keyboard, and revise with a written evidence statement.",
        "rubric_title": "Chromatic Analysis-Composition Rubric",
        "rubric_note": "Use for modulation analyses, chromatic-harmony writing, colloquia, and the analysis-composition capstone.",
        "rubric": [
            ("Chromatic accuracy", "Altered tones, chord identity, inversion, and resolution are consistently correct.", "Minor slips do not change function.", "Several spelling or resolution errors weaken control.", "Errors prevent a coherent chromatic reading."),
            ("Functional/formal reasoning", "Relates local voice leading to key, phrase, and form with strong evidence.", "Reasoning is sound with small gaps.", "Reasoning is partial or inconsistent.", "Functional claims are unsupported."),
            ("Compositional control", "Chromaticism has structural purpose and persuasive musical effect.", "Chromaticism is functional and mostly controlled.", "Effect is uneven or mechanically applied.", "Chromatic choices lack resolution or purpose."),
            ("Communication and revision", "Score, map, defense, and revision are precise and professional.", "Materials are clear and revisions mostly succeed.", "Clarity or revision is inconsistent.", "Materials are incomplete or unclear."),
        ],
    },
    "Sight Singing": {
        "filename": "Sight_Singing_Syllabus_and_Lesson_Plans.docx",
        "code": "[Insert course code]",
        "text": "The Sight Singer, Volume 1, Audrey Snyder",
        "meeting_model": "Suggested model: two 60-minute laboratory meetings weekly; adjust to the approved institutional schedule.",
        "meeting_names": ["Rhythm and tonal preparation", "Read, perform, and reflect"],
        "visuals": [
            ("V7_sight_singing_routine.png", "Sight-Singing Preparation Routine", "A scan-audiate-perform-diagnose-retry routine for first-read independence."),
            ("V2_rhythm_tree.png", "Rhythm Value Tree", "A hierarchy showing beat division and common rhythmic relationships."),
        ],
        "support": "Provide a tonic drone, slower tempo, chunked rhythm or tonal cells, hand signs, an octave option, and a private or small-group first attempt.",
        "extension": "Reduce preparation time, begin on a non-tonic degree, remove the drone, add a partner line, transpose, or require a one-take performance.",
        "practice": "Practice 10-15 minutes on the assigned Snyder exercises: preview silently, record or perform one take, diagnose one issue, apply one strategy, and complete a second take.",
        "rubric_title": "Sight-Singing Performance Rubric",
        "rubric_note": "Score the first read and any coached retry separately when growth evidence is required. Protect vocal health and offer an appropriate octave.",
        "rubric": [
            ("Preparation", "Establishes tonic, meter, pulse, starting pitch, and hazards independently.", "Establishes most elements with minimal prompting.", "Preparation is incomplete or needs repeated cues.", "Begins without a usable tonal or metric plan."),
            ("Pitch/intonation", "Pitch and tonal relationships are consistently secure.", "Mostly accurate with quick recovery.", "Recurring pitch errors interrupt tonal orientation.", "Pitch center is not maintained."),
            ("Rhythm/continuity", "Pulse, subdivision, and continuity remain secure throughout.", "Minor errors occur without stopping.", "Pulse or continuity breaks more than once.", "Reading repeatedly stops or loses meter."),
            ("Musicianship/recovery", "Phrasing is musical and recovery uses clear anchor tones.", "Musical shape and recovery are generally effective.", "Expression or recovery is inconsistent.", "No effective recovery strategy is evident."),
        ],
    },
}


INSTITUTIONAL_FIELDS = [
    ("Institution", "[Insert college/university]"), ("College/Department", "[Insert unit]"),
    ("Instructor", "[Insert name and credentials]"), ("Academic term", "[Insert term and AY]"),
    ("Class schedule", "[Insert days/time/room]"), ("Consultation", "[Insert office hours/contact]"),
]


def set_document_context(new_doc):
    global doc, BULLET_ID, NUMBER_ID
    doc = new_doc
    sec = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = False
    sec.different_first_page_header_footer = False
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    configure_styles(doc)
    BULLET_ID, NUMBER_ID = add_numbering(doc)
    return sec


def add_cover(course_name, desc, cfg, units):
    add_header_footer(doc.sections[0], course_name.upper(), "SYLLABUS + 18-WEEK LESSON PLANS")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("COLLEGE MUSIC TEACHING PACKAGE")
    set_font(r, size=9, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(course_name)
    set_font(r, size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Course Syllabus, Weekly Learning Map, Activities, Rubric, Visual Aids, and Detailed Lesson Plans")
    set_font(r, size=13, italic=True, color=TEAL)
    add_callout("Course purpose:", desc)
    add_heading("Course-at-a-glance", 2)
    add_labeled_paragraph("Recommended load:", units, after=3)
    add_labeled_paragraph("Primary text:", cfg["text"], after=3)
    add_labeled_paragraph("Planning model:", cfg["meeting_model"], after=3)
    add_labeled_paragraph("Edition note:", "Insert exact chapter, unit, exercise, and page numbers from the edition owned by the class; page numbering varies by printing.", after=10)
    add_heading("Before releasing to students", 2)
    for item in [
        "Replace all bracketed institutional fields and publish the official course code and contact hours.",
        "Align grading equivalents, attendance rules, and make-up procedures with current university policy.",
        "Confirm room, piano/keyboard, projection, LMS, printing, and approved recording arrangements.",
        "Choose examples that are public-domain, licensed, teacher-composed, or institutionally cleared.",
    ]:
        add_bullet(item)


def add_course_setup(course_name, prereq, cfg, units):
    page_break()
    add_section_title(f"{course_name}: Course Setup", "Complete the fields below before circulation.", "Instructor Setup")
    add_heading("Institutional information", 2)
    for label, value in INSTITUTIONAL_FIELDS:
        add_labeled_paragraph(f"{label}:", value, after=3)
    add_heading("Catalog and delivery profile", 2)
    add_labeled_paragraph("Course code:", cfg["code"])
    add_labeled_paragraph("Units/contact hours:", units)
    add_labeled_paragraph("Prerequisite:", prereq)
    add_labeled_paragraph("Primary text:", cfg["text"])
    add_labeled_paragraph("Delivery mode:", "In-person, blended, or online as approved by the institution.")
    add_labeled_paragraph("Learning platforms:", "[Insert LMS, notation software, recording platform, or approved alternatives].")
    add_heading("Required and recommended materials", 2)
    materials = [
        "Assigned textbook and staff paper; pencil, eraser, and folder or digital portfolio.",
        "Access to a piano/keyboard or approved virtual keyboard for pitch and harmony checks.",
        "Headphones and an institution-approved recording method when an assessment requires audio.",
        "Instructor-provided excerpts, worksheets, and visual aids in accessible print/digital formats.",
    ]
    if course_name == "Sight Singing":
        materials[1] = "Pitch reference (piano/keyboard or approved app), water, and a quiet practice space; students may sing in an appropriate octave."
    for item in materials:
        add_bullet(item)


def add_outcomes_and_assessment(course_name, outcomes, assessments):
    page_break()
    add_section_title(f"{course_name}: Outcomes and Assessment", "Outcome-aligned evidence with a built-in cycle of attempt, feedback, correction, and demonstration.", "Syllabus Core")
    add_heading("Course learning outcomes", 1)
    add_course_outcomes(outcomes)
    add_heading("Assessment weights", 1)
    add_assessment_table(assessments)
    add_callout("Grade calculation:", "Calculate the weighted raw score above, then apply the official university grading or transmutation system. Publish the approved equivalent table in the LMS or course handout.")
    add_heading("Evidence cycle", 2)
    for step in [
        "Prepare through a short reading, listening, writing, keyboard, or vocal task.",
        "Attempt a low-stakes first try that reveals current understanding.",
        "Diagnose the specific musical issue using shared vocabulary.",
        "Coach and revise with one targeted strategy and timely feedback.",
        "Demonstrate independently, then archive the evidence in the course portfolio.",
    ]:
        add_bullet(step)


def add_policies(course_name):
    page_break()
    add_section_title(f"{course_name}: Learning and Course Policies", "Adapt the bracketed items to official institutional policy; never let this section override university rules.", "Student Guide")
    policies = [
        ("Attendance and participation", "Music literacy develops through cumulative guided practice. Follow the institution's approved attendance policy. Students should notify the instructor through the official channel and complete the announced recovery task after an absence."),
        ("Late and make-up work", "Insert the department-approved window, documentation requirements, and procedure for quizzes, examinations, performances, and technical disruptions. Equivalent make-up tasks may use different musical material."),
        ("Academic integrity", "Submitted analysis, notation, composition, recordings, and reflections must represent the student's own work except where collaboration is explicitly assigned. Credit every borrowed score, recording, image, or idea in the format required by the instructor."),
        ("Generative AI and automated tools", "AI, notation playback, tuners, analyzers, and answer-generating tools may be used only for the purposes explicitly permitted on an assignment. Students must disclose permitted assistance and remain able to perform, explain, and revise the work independently. Undisclosed or prohibited assistance is handled under institutional academic-integrity policy."),
        ("Accessibility and inclusive participation", "Students may request reasonable accommodations through the authorized campus office. Provide accessible copies, clear notation, usable contrast, flexible demonstration modes when outcomes allow, and an appropriate vocal octave or range. Accommodation does not lower the stated musical outcome."),
        ("Recording and privacy", "Record students only when institutionally authorized and educationally necessary. State the purpose, storage location, access, retention period, and non-recording alternative. Do not post class recordings publicly without valid consent and approval."),
        ("Copyright and repertoire", "Use lawfully acquired textbooks and public-domain, licensed, teacher-composed, or institutionally cleared excerpts. Do not distribute scans or recordings beyond the permitted educational use."),
        ("Respectful musicianship", "Critique the musical evidence, not the person. Use names and pronouns correctly, support developing voices and instruments, and maintain a rehearsal environment where students can attempt, recover, and improve safely."),
    ]
    for title, body in policies:
        add_heading(title, 2)
        add_labeled_paragraph("Policy:", body)


def split_activity_text(text):
    parts = [part.strip() for part in text.split(";") if part.strip()]
    while len(parts) < 3:
        parts.append(parts[-1] if parts else "guided practice")
    return parts


def add_forced_section_title(title, subtitle=None, kicker=None):
    """Start a section on a new page without an empty page-break paragraph."""
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(kicker.upper())
        set_font(r, size=9, bold=True, color=GOLD)
    p = doc.add_paragraph()
    if not kicker:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=24, bold=True, color=NAVY)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        sr = sp.add_run(subtitle)
        set_font(sr, size=11.5, italic=True, color=TEAL)


def lesson_objectives(focus, target, evidence):
    return [
        f"Identify and explain the central features of {focus.lower()}.",
        f"Demonstrate this learning target: {target}.",
        f"Produce, perform, or submit the announced evidence: {evidence}.",
    ]


def add_lesson_plan(course_name, week, focus, target, activities, evidence, cfg):
    page_break()
    add_section_title(f"Week {week} Lesson Plan", focus, course_name)
    add_callout("Weekly target:", target)
    add_labeled_paragraph("Textbook focus:", f"{cfg['text']} - assign the matching unit, exercise, and page numbers from the class edition.", after=3)
    add_labeled_paragraph("Materials / visual aids:", f"Textbook, staff paper or score, board/projector, piano or keyboard, playback/recording only when approved; use the visual-aid appendix and the weekly visual code when named.", after=6)
    add_heading("Measurable objectives", 2)
    for objective in lesson_objectives(focus, target, evidence):
        add_bullet(objective)

    activity_parts = split_activity_text(activities)
    meeting_names = cfg["meeting_names"]
    if len(meeting_names) == 3:
        flows = [
            "5 min retrieval; 8 min aural/visual hook; 12 min concise model; 20 min guided practice; 10 min independent check; 5 min exit ticket.",
            "5 min warm-up; 10 min error review; 20 min coached application; 15 min transfer task; 7 min peer feedback; 3 min reflection.",
            "5 min retrieval; 10 min misconception mini-lesson; 20 min synthesis or creation; 15 min assessed evidence; 7 min correction/retry; 3 min assignment briefing.",
        ]
        products = [
            "Annotated example or accurately modeled first attempt.",
            "Completed application with peer/instructor feedback marked.",
            f"Independent evidence: {evidence}.",
        ]
    else:
        flows = [
            "8 min physical/vocal setup; 12 min pulse and rhythm work; 12 min tonal orientation; 18 min guided reading; 7 min partner feedback; 3 min reflection.",
            "5 min silent scan; 10 min rhythm/tonal preparation; 20 min individual and ensemble first reads; 12 min diagnosis and coached retry; 8 min assessment; 5 min practice plan.",
        ]
        products = [
            "Preview checklist plus accurately prepared rhythm/tonal cells.",
            f"First read, coached retry, and evidence: {evidence}.",
        ]

    for idx, name in enumerate(meeting_names):
        add_heading(f"Meeting {chr(65 + idx)} - {name}", 2)
        add_labeled_paragraph("Focus activity:", activity_parts[min(idx, len(activity_parts) - 1)] + ".", after=2)
        add_labeled_paragraph("Flow:", flows[idx], after=2)
        add_labeled_paragraph("Student evidence:", products[idx], after=4)

    add_heading("Assessment, differentiation, and follow-through", 2)
    add_labeled_paragraph("Formative checks:", "Cold/warm response, one-minute construction or performance, error annotation, partner explanation, and exit ticket. Record only the evidence needed for the stated outcome.", after=2)
    add_labeled_paragraph("Support:", cfg["support"], after=2)
    add_labeled_paragraph("Extension:", cfg["extension"], after=2)
    add_labeled_paragraph("Assignment / practice:", cfg["practice"], after=2)
    add_labeled_paragraph("Instructor reflection:", "What evidence showed learning? Which misconception or barrier needs reteaching? What will change in the next meeting? ______________________________", after=2)


def add_visual_appendix(course_name, cfg):
    page_break()
    add_section_title(f"{course_name}: Visual-Aid Appendix", "Project, print, or post these aids in an accessible format. Pair every diagram with spoken and written explanation.", "Teaching Resources")
    for index, (filename, title, alt_text) in enumerate(cfg["visuals"]):
        if index and index % 2 == 0:
            add_forced_section_title(f"{course_name}: Visual-Aid Appendix (continued)", None, "Teaching Resources")
        add_heading(title, 2)
        path = ASSET_DIR / filename
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(5.65))
        inline = doc.inline_shapes[-1]._inline
        inline.docPr.set("title", title)
        inline.docPr.set("descr", alt_text)
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        r = caption.add_run(f"{title}. {alt_text}")
        set_font(r, size=8.5, italic=True, color=GRAY)


def add_rubric_and_checklist(course_name, cfg):
    add_forced_section_title(f"{course_name}: Rubric and Implementation Checklist", "Publish task-specific scoring priorities before students begin.", "Assessment Appendix")
    add_rubric(cfg["rubric_title"], cfg["rubric"], cfg["rubric_note"])
    add_heading("Instructor implementation checklist", 1)
    for item in [
        "Insert the official course code, term, section, schedule, room, consultation hours, and contact channel.",
        "Map each week to exact textbook units/exercises/pages and the approved academic calendar, including holidays and examination dates.",
        "Publish the university's current grading equivalents, attendance, make-up, integrity, accessibility, privacy, and emergency procedures.",
        "Prepare accessible visual aids, lawful excerpts, answer keys, differentiated examples, and a plan for students without personal devices.",
        "Calibrate rubric examples, moderation procedures, feedback turnaround, and secure storage of grades and recordings.",
        "After each week, record evidence, reteaching needs, pacing changes, and materials to replace before the next offering.",
    ]:
        add_bullet(item)
    add_heading("Instructor approval", 2)
    add_labeled_paragraph("Prepared by:", "____________________________________    Date: ____________________")
    add_labeled_paragraph("Reviewed by:", "____________________________________    Date: ____________________")
    add_labeled_paragraph("Approved by:", "____________________________________    Date: ____________________")


def build_course(course_tuple):
    course_name, desc, prereq, units, outcomes, assessments, weeks, activities = course_tuple
    cfg = COURSE_CONFIG[course_name]
    new_doc = Document()
    set_document_context(new_doc)
    add_cover(course_name, desc, cfg, units)
    add_course_setup(course_name, prereq, cfg, units)
    add_outcomes_and_assessment(course_name, outcomes, assessments)
    add_policies(course_name)
    add_weekly_schedule(course_name, weeks)
    add_signature_activities(course_name, activities)
    for week_data in weeks:
        add_lesson_plan(course_name, *week_data, cfg)
    add_visual_appendix(course_name, cfg)
    add_rubric_and_checklist(course_name, cfg)

    props = new_doc.core_properties
    props.title = f"{course_name} Syllabus and 18-Week Lesson Plans"
    props.subject = "College-level music curriculum for the Philippines"
    props.author = "Course Instructor"
    props.keywords = "music syllabus, lesson plans, college, Philippines"
    output = ROOT / cfg["filename"]
    new_doc.save(output)
    return output


if __name__ == "__main__":
    outputs = [build_course(course) for course in courses]
    for output in outputs:
        print(output)
